"""Core document, style-analysis, and content-generation services.

The module is deliberately UI-agnostic so the same behavior can be used from
Streamlit, tests, a future API, or a background worker.
"""

from __future__ import annotations

import io
import json
import math
import re
from collections import Counter
from collections.abc import Iterable
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, BinaryIO, Protocol

MAX_DOCUMENT_CHARS = 240_000


class AgentError(RuntimeError):
    """A safe, user-facing error raised by an AI service."""


class UploadedFile(Protocol):
    name: str

    def read(self, size: int = -1) -> bytes: ...


@dataclass(slots=True)
class ProcessedDocument:
    name: str
    text: str
    file_type: str
    word_count: int
    character_count: int


@dataclass(slots=True)
class StyleProfile:
    tone: str
    voice: str
    structure: str
    vocabulary: str
    sentence_patterns: str
    examples_style: str
    signature_traits: list[str] = field(default_factory=list)
    top_terms: list[str] = field(default_factory=list)
    readability_score: float = 0.0
    average_sentence_length: float = 0.0
    source_word_count: int = 0
    analysis_mode: str = "local"
    analyzed_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, value: dict[str, Any]) -> StyleProfile:
        allowed = cls.__dataclass_fields__.keys()
        return cls(**{key: value[key] for key in allowed if key in value})

    def compact_prompt(self) -> str:
        traits = ", ".join(self.signature_traits) or "clear, reader-focused prose"
        terms = ", ".join(self.top_terms[:8]) or "topic-specific language"
        return (
            f"Tone: {self.tone}\n"
            f"Voice: {self.voice}\n"
            f"Structure: {self.structure}\n"
            f"Vocabulary: {self.vocabulary}\n"
            f"Sentence rhythm: {self.sentence_patterns}\n"
            f"Examples: {self.examples_style}\n"
            f"Signature traits: {traits}\n"
            f"Recurring terms: {terms}"
        )


@dataclass(slots=True)
class GenerationRequest:
    topic: str
    audience: str = "Curious professionals"
    article_type: str = "How-to guide"
    length: str = "Medium · 800–1,000 words"
    keywords: list[str] = field(default_factory=list)
    instructions: str = ""


@dataclass(slots=True)
class GeneratedArticle:
    title: str
    content: str
    excerpt: str
    word_count: int
    mode: str
    created_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())


class DocumentProcessor:
    """Extract text locally from TXT, Markdown, PDF, and DOCX files."""

    supported_extensions = {".txt", ".md", ".pdf", ".docx"}

    def process(self, uploaded: UploadedFile | BinaryIO, name: str | None = None) -> ProcessedDocument:
        filename = name or getattr(uploaded, "name", "document.txt")
        suffix = Path(filename).suffix.lower()
        if suffix not in self.supported_extensions:
            raise ValueError(
                f"Unsupported file type '{suffix or 'unknown'}'. Upload PDF, DOCX, TXT, or Markdown."
            )

        data = uploaded.read()
        if isinstance(data, str):
            data = data.encode("utf-8")
        if not data:
            raise ValueError(f"{filename} is empty.")

        if suffix in {".txt", ".md"}:
            text = self._decode_text(data)
        elif suffix == ".pdf":
            text = self._extract_pdf(data)
        else:
            text = self._extract_docx(data)

        text = self._clean_text(text)[:MAX_DOCUMENT_CHARS]
        if len(text.split()) < 25:
            raise ValueError(f"{filename} does not contain enough readable text (minimum 25 words).")
        return ProcessedDocument(
            name=filename,
            text=text,
            file_type=suffix.removeprefix(".").upper(),
            word_count=len(re.findall(r"\b[\w’'-]+\b", text)),
            character_count=len(text),
        )

    @staticmethod
    def _decode_text(data: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252", "latin-1"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="replace")

    @staticmethod
    def _extract_pdf(data: bytes) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(data))
            return "\n\n".join((page.extract_text() or "") for page in reader.pages)
        except ImportError as exc:
            raise ValueError("PDF support is unavailable. Install the pypdf dependency.") from exc
        except Exception as exc:
            raise ValueError("The PDF could not be read. It may be scanned, encrypted, or damaged.") from exc

    @staticmethod
    def _extract_docx(data: bytes) -> str:
        try:
            from docx import Document

            document = Document(io.BytesIO(data))
            blocks = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n\n".join(blocks)
        except ImportError as exc:
            raise ValueError("DOCX support is unavailable. Install the python-docx dependency.") from exc
        except Exception as exc:
            raise ValueError(
                "The DOCX file could not be read. It may be damaged or password-protected."
            ) from exc

    @staticmethod
    def _clean_text(text: str) -> str:
        text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


class DigitalOceanAgent:
    """Small adapter for a DigitalOcean Gradient AI agent endpoint."""

    def __init__(
        self,
        endpoint: str | None,
        access_key: str | None,
        *,
        entity_id: str = "anonymous-writer",
        memori_api_key: str | None = None,
    ) -> None:
        self.endpoint = (endpoint or "").strip().rstrip("/")
        self.access_key = (access_key or "").strip()
        self.entity_id = entity_id
        self.memori_api_key = (memori_api_key or "").strip()
        self.client: Any | None = None
        self.memory_enabled = False
        self.memory_error: str | None = None

        if self.configured:
            self.client = self._build_client()

    @property
    def configured(self) -> bool:
        return bool(self.endpoint and self.access_key)

    @property
    def base_url(self) -> str:
        if self.endpoint.endswith("/api/v1"):
            return f"{self.endpoint}/"
        return f"{self.endpoint}/api/v1/"

    def _build_client(self) -> Any:
        try:
            from openai import OpenAI
        except ImportError as exc:
            raise AgentError("The OpenAI compatibility client is not installed.") from exc

        client = OpenAI(base_url=self.base_url, api_key=self.access_key, timeout=90.0)
        if self.memori_api_key:
            try:
                # Memori reads MEMORI_API_KEY from its standard environment/config.
                # Passing it explicitly is supported by current Cloud clients and the
                # fallback keeps older SDKs compatible.
                from memori import Memori

                try:
                    memory = Memori(api_key=self.memori_api_key).llm.register(client)
                except TypeError:
                    memory = Memori().llm.register(client)
                memory.attribution(entity_id=self.entity_id, process_id="blog-writing-agent")
                self.memory_enabled = True
            except Exception as exc:  # Memory should never make writing unavailable.
                self.memory_error = str(exc)
        return client

    def complete(
        self,
        messages: list[dict[str, str]],
        *,
        temperature: float = 0.7,
        max_tokens: int = 2400,
    ) -> str:
        if not self.client:
            raise AgentError("DigitalOcean AI is not configured.")
        try:
            response = self.client.chat.completions.create(
                model="n/a",
                messages=messages,
                temperature=temperature,
                max_tokens=max_tokens,
            )
            content = response.choices[0].message.content
            if not content:
                raise AgentError("The AI returned an empty response. Please try again.")
            return content.strip()
        except AgentError:
            raise
        except Exception as exc:
            message = str(exc)
            if "401" in message or "authentication" in message.lower():
                raise AgentError("DigitalOcean rejected the access key. Check your credentials.") from exc
            if "timeout" in message.lower():
                raise AgentError("DigitalOcean took too long to respond. Please retry.") from exc
            raise AgentError(f"DigitalOcean AI request failed: {message[:180]}") from exc


class StyleAnalyzer:
    """Create a rich style fingerprint using AI or local linguistic heuristics."""

    STOP_WORDS = {
        "about",
        "after",
        "again",
        "also",
        "among",
        "because",
        "been",
        "before",
        "being",
        "between",
        "both",
        "could",
        "does",
        "each",
        "from",
        "have",
        "into",
        "just",
        "more",
        "most",
        "much",
        "only",
        "other",
        "over",
        "same",
        "some",
        "such",
        "than",
        "that",
        "their",
        "them",
        "then",
        "there",
        "these",
        "they",
        "this",
        "those",
        "through",
        "very",
        "what",
        "when",
        "where",
        "which",
        "while",
        "with",
        "would",
        "your",
        "you",
        "were",
        "will",
        "shall",
        "the",
        "and",
        "for",
        "are",
        "but",
        "not",
        "can",
        "our",
        "its",
        "has",
        "had",
        "was",
    }

    def __init__(self, agent: DigitalOceanAgent | None = None) -> None:
        self.agent = agent

    def analyze(self, documents: Iterable[ProcessedDocument]) -> StyleProfile:
        docs = list(documents)
        if not docs:
            raise ValueError("Upload at least one document before analyzing your style.")
        text = "\n\n--- ARTICLE BREAK ---\n\n".join(doc.text for doc in docs)
        if self.agent and self.agent.configured:
            try:
                return self._analyze_with_ai(text)
            except AgentError:
                raise
            except Exception as exc:
                raise AgentError(f"The style profile could not be parsed: {exc}") from exc
        return self._analyze_locally(text)

    def _analyze_with_ai(self, text: str) -> StyleProfile:
        sample = text[:80_000]
        system = (
            "You are an expert editorial stylist. Analyze the author's repeatable writing "
            "habits, not the article subject. Return strict JSON only with these keys: tone, "
            "voice, structure, vocabulary, sentence_patterns, examples_style, signature_traits "
            "(array of 3-6 short strings), top_terms (array of up to 8 meaningful terms), "
            "readability_score (0-100), average_sentence_length (number). Be specific, concise, "
            "constructive, and never copy long passages from the source."
        )
        raw = self.agent.complete(
            [
                {"role": "system", "content": system},
                {"role": "user", "content": f"Analyze this writing sample:\n\n{sample}"},
            ],
            temperature=0.2,
            max_tokens=1100,
        )
        payload = _extract_json(raw)
        payload.update(
            source_word_count=len(re.findall(r"\b[\w’'-]+\b", text)),
            analysis_mode="DigitalOcean AI + Memori" if self.agent.memory_enabled else "DigitalOcean AI",
        )
        return StyleProfile.from_dict(payload)

    def _analyze_locally(self, text: str) -> StyleProfile:
        words = re.findall(r"\b[A-Za-z][A-Za-z’'-]*\b", text)
        lowered = [word.lower().replace("’", "'") for word in words]
        sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if len(s.split()) > 2]
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        sentence_lengths = [len(re.findall(r"\b[\w’'-]+\b", s)) for s in sentences] or [len(words)]
        average_sentence = sum(sentence_lengths) / max(len(sentence_lengths), 1)
        variance = sum((n - average_sentence) ** 2 for n in sentence_lengths) / max(len(sentence_lengths), 1)
        rhythm_spread = math.sqrt(variance)
        contractions = sum(1 for word in lowered if "'" in word)
        questions = text.count("?")
        exclamations = text.count("!")
        second_person = sum(lowered.count(p) for p in ("you", "your", "yours"))
        first_person = sum(lowered.count(p) for p in ("i", "we", "my", "our", "us"))
        formal_markers = sum(
            lowered.count(word)
            for word in ("therefore", "however", "moreover", "consequently", "furthermore", "nevertheless")
        )

        if contractions + exclamations > max(len(words) / 90, 3):
            tone = "Warm and conversational, with an easygoing, energetic edge"
        elif formal_markers > max(len(sentences) / 15, 2) or average_sentence > 23:
            tone = "Polished and analytical, with a measured professional tone"
        else:
            tone = "Clear and approachable, balancing confidence with friendliness"

        if second_person > first_person * 1.3 and second_person > 3:
            voice = "Reader-direct and coaching; frequently addresses the audience as ‘you’"
        elif first_person > second_person and first_person > 3:
            voice = "Personal and experience-led, using first-person perspective to build trust"
        else:
            voice = "Confident explanatory voice with an objective, helpful point of view"

        avg_paragraph = len(words) / max(len(paragraphs), 1)
        if avg_paragraph < 55:
            structure = "Scannable, compact paragraphs with frequent idea breaks and clear progression"
        elif avg_paragraph > 110:
            structure = "Long-form, developed paragraphs that build arguments in depth"
        else:
            structure = "Moderate paragraphs that move from context to explanation to takeaway"

        unique_ratio = len(set(lowered)) / max(len(lowered), 1)
        long_word_ratio = sum(len(word) >= 9 for word in lowered) / max(len(lowered), 1)
        if long_word_ratio > 0.16:
            vocabulary = "Advanced and domain-aware, with precise terminology and varied word choice"
        elif unique_ratio > 0.48:
            vocabulary = "Varied but accessible, favoring vivid words over unnecessary jargon"
        else:
            vocabulary = "Plainspoken and accessible, prioritizing clarity and familiar language"

        if rhythm_spread > 10:
            patterns = f"Highly varied rhythm; sentences average {average_sentence:.1f} words with deliberate short-long contrast"
        elif average_sentence < 15:
            patterns = f"Crisp, fast-moving sentences averaging {average_sentence:.1f} words"
        elif average_sentence > 24:
            patterns = f"Layered, complex sentences averaging {average_sentence:.1f} words"
        else:
            patterns = f"Balanced sentence rhythm averaging {average_sentence:.1f} words"

        example_markers = len(
            re.findall(r"\b(for example|for instance|such as|imagine|consider)\b", text, re.I)
        )
        if example_markers >= max(len(sentences) / 18, 2):
            examples_style = "Frequently grounds ideas in concrete examples, scenarios, and analogies"
        else:
            examples_style = (
                "Uses examples selectively, leaning on direct explanation and practical takeaways"
            )

        traits: list[str] = []
        if questions >= max(len(sentences) / 18, 2):
            traits.append("Uses questions to pull readers forward")
        if second_person > 4:
            traits.append("Speaks directly to the reader")
        if exclamations > 1:
            traits.append("Adds occasional bursts of enthusiasm")
        if rhythm_spread > 8:
            traits.append("Mixes punchy lines with developed explanations")
        if len(paragraphs) > 3 and avg_paragraph < 80:
            traits.append("Keeps paragraphs deliberately scannable")
        if formal_markers > 1:
            traits.append("Uses explicit transitions to guide the argument")
        if example_markers > 1:
            traits.append("Makes abstract ideas concrete")
        traits.extend(["Leads with clarity", "Ends sections with a usable takeaway"])
        traits = list(dict.fromkeys(traits))[:6]

        counts = Counter(word for word in lowered if len(word) > 4 and word not in self.STOP_WORDS)
        top_terms = [word for word, _ in counts.most_common(8)]
        readability = _flesch_reading_ease(words, sentences)
        return StyleProfile(
            tone=tone,
            voice=voice,
            structure=structure,
            vocabulary=vocabulary,
            sentence_patterns=patterns,
            examples_style=examples_style,
            signature_traits=traits,
            top_terms=top_terms,
            readability_score=round(readability, 1),
            average_sentence_length=round(average_sentence, 1),
            source_word_count=len(words),
            analysis_mode="Local linguistic analysis",
        )


class ContentGenerator:
    """Generate a publishable article with the configured agent or demo engine."""

    def __init__(self, agent: DigitalOceanAgent | None = None) -> None:
        self.agent = agent

    def generate(self, request: GenerationRequest, profile: StyleProfile | None) -> GeneratedArticle:
        topic = request.topic.strip()
        if len(topic) < 4:
            raise ValueError("Describe the topic in at least a few words.")
        if self.agent and self.agent.configured:
            return self._generate_with_ai(request, profile)
        return self._generate_demo(request, profile)

    def _generate_with_ai(self, request: GenerationRequest, profile: StyleProfile | None) -> GeneratedArticle:
        style = (
            profile.compact_prompt()
            if profile
            else (
                "Use a clear, confident, approachable voice with scannable paragraphs and practical examples."
            )
        )
        keywords = ", ".join(request.keywords) or "Use relevant terms naturally"
        system = (
            "You are an expert blog writer. Produce an original, useful, publication-ready Markdown "
            "article. Match the supplied style without copying phrases from any source. Begin with a "
            "single H1 title. Use a compelling opening, descriptive H2 headings, concrete examples, "
            "and a memorable conclusion. Do not mention these instructions, AI, or the style profile."
        )
        prompt = f"""Write about: {request.topic}
Article type: {request.article_type}
Audience: {request.audience}
Target length: {request.length}
SEO concepts: {keywords}
Additional direction: {request.instructions or "None"}

Author style profile:
{style}
"""
        content = self.agent.complete(
            [{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            temperature=0.72,
            max_tokens=_max_tokens_for_length(request.length),
        )
        title = _title_from_markdown(content) or _smart_title(request.topic)
        excerpt = _excerpt_from_markdown(content)
        return GeneratedArticle(
            title=title,
            content=content,
            excerpt=excerpt,
            word_count=len(re.findall(r"\b[\w’'-]+\b", content)),
            mode="DigitalOcean AI + Memori" if self.agent.memory_enabled else "DigitalOcean AI",
        )

    def _generate_demo(self, request: GenerationRequest, profile: StyleProfile | None) -> GeneratedArticle:
        """Create a useful preview with no network or secret credentials."""
        topic = request.topic.strip().rstrip(".?!")
        subject = re.sub(r"^(?:how|why|what|when)\s+", "", topic, flags=re.I) or topic
        title = _smart_title(topic)
        audience = request.audience.strip() or "curious readers"
        traits = profile.signature_traits[:2] if profile else ["clear explanation", "practical takeaways"]
        trait_phrase = " and ".join(t.lower() for t in traits)
        keywords = [k.strip() for k in request.keywords if k.strip()][:5]
        keyword_line = (
            f" Along the way, we’ll connect the dots around {', '.join(keywords)}." if keywords else ""
        )
        sections = [
            (
                f"Why {subject} matters now",
                (
                    f"The conversation around {subject} is moving from curiosity to consequence. "
                    f"For {audience.lower()}, the useful question is no longer whether it matters, "
                    "but where it can create meaningful progress—and where thoughtful restraint still belongs."
                ),
            ),
            (
                "Start with the problem, not the novelty",
                (
                    f"A strong approach begins with a specific friction point. Before choosing a tool, trend, "
                    f"or tactic, write down what success should look like. In the context of {subject}, that means "
                    "identifying the people affected, the repeated task that needs improvement, and the signal "
                    "that will tell you the change is working."
                ),
            ),
            (
                "Turn the idea into a repeatable system",
                (
                    "Small experiments beat sweeping transformations. Choose one workflow, establish a baseline, "
                    "and run a short pilot. Keep the feedback loop visible: observe what happened, capture what "
                    "surprised you, and adjust one variable at a time. That discipline turns an exciting idea "
                    "into something a team can trust."
                ),
            ),
            (
                "Keep the human judgment in the loop",
                (
                    f"The best results rarely come from automation alone. They come from pairing consistent systems "
                    f"with context, taste, and accountability. Treat {subject} as a collaborator in the process—not "
                    "a substitute for the people who understand the audience, the risks, and the goal."
                ),
            ),
        ]
        content_parts = [
            f"# {title}",
            (
                f"{topic.capitalize()} can feel like one of those ideas everyone is discussing and few are "
                f"turning into a durable practice. This guide makes it concrete. We’ll focus on {trait_phrase}, "
                f"so you can move from interest to a plan that holds up in the real world.{keyword_line}"
            ),
        ]
        for heading, body in sections:
            content_parts.extend([f"## {heading}", body])
            if heading == "Turn the idea into a repeatable system":
                content_parts.append(
                    "1. **Define one outcome.** Make it observable and relevant.\n"
                    "2. **Run a contained pilot.** Learn before you scale.\n"
                    "3. **Review the evidence.** Include the people doing the work.\n"
                    "4. **Document the pattern.** Make the next attempt easier than the first."
                )
        content_parts.extend(
            [
                "## A practical next step",
                (
                    f"Pick one decision or workflow connected to {subject} and spend 20 minutes mapping how it works "
                    "today. Circle the slowest or least clear moment. That is your starting point—not a grand "
                    "transformation, just one honest opportunity to make the work better."
                ),
                (
                    "Progress becomes sustainable when the next move is small enough to begin and meaningful enough "
                    "to measure. Start there, learn quickly, and let the evidence shape what comes next."
                ),
            ]
        )
        content = "\n\n".join(content_parts)
        if request.instructions.strip():
            content += f"\n\n> **Editorial note:** {request.instructions.strip()}"
        excerpt = f"A practical, people-first guide to turning {topic} from an interesting idea into a repeatable, measurable practice."
        return GeneratedArticle(
            title=title,
            content=content,
            excerpt=excerpt,
            word_count=len(re.findall(r"\b[\w’'-]+\b", content)),
            mode="Local demo engine",
        )


def _extract_json(raw: str) -> dict[str, Any]:
    cleaned = raw.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    start, end = cleaned.find("{"), cleaned.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("AI response did not contain a JSON object")
    return json.loads(cleaned[start : end + 1])


def _max_tokens_for_length(length: str) -> int:
    lower = length.lower()
    if "short" in lower or "500" in lower:
        return 1400
    if "long" in lower or "1,500" in lower or "1500" in lower:
        return 3600
    return 2500


def _smart_title(topic: str) -> str:
    clean = re.sub(r"\s+", " ", topic.strip())
    if clean.lower().startswith(("how ", "why ", "what ", "when ")):
        return clean[0].upper() + clean[1:]
    return f"A Practical Guide to {clean.title()}"


def _title_from_markdown(content: str) -> str:
    match = re.search(r"^#\s+(.+)$", content, re.M)
    return match.group(1).strip() if match else ""


def _excerpt_from_markdown(content: str) -> str:
    paragraphs = [
        re.sub(r"[*_>`#]", "", p).strip()
        for p in re.split(r"\n\s*\n", content)
        if p.strip() and not p.lstrip().startswith("#")
    ]
    excerpt = paragraphs[0] if paragraphs else content
    return excerpt[:237].rstrip() + ("…" if len(excerpt) > 237 else "")


def _flesch_reading_ease(words: list[str], sentences: list[str]) -> float:
    if not words or not sentences:
        return 0.0
    syllables = sum(_estimate_syllables(word) for word in words)
    score = 206.835 - 1.015 * (len(words) / len(sentences)) - 84.6 * (syllables / len(words))
    return min(100.0, max(0.0, score))


def _estimate_syllables(word: str) -> int:
    clean = re.sub(r"[^a-z]", "", word.lower())
    if not clean:
        return 1
    groups = len(re.findall(r"[aeiouy]+", clean))
    if clean.endswith("e") and not clean.endswith(("le", "ye")) and groups > 1:
        groups -= 1
    return max(1, groups)
